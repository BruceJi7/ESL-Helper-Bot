import docx
from os import path
from datetime import datetime



def writeMessageListToDoc(listOfMessages, title):

    doc = docx.Document()

    dateString = datetime.strftime(datetime.today(), '%Y-%m-%d')
    
    # Write messages to file
    doc.add_heading(f"Session notes from {dateString}", 2)
    doc.add_paragraph()
    for m in listOfMessages:
        doc.add_paragraph(m)

    # Handle filename/path
    basePath = path.dirname(path.abspath(__file__))
    docx_filename_path = path.join(basePath, f'{title}.docx')

    # Handle if file already exists
    if path.exists(docx_filename_path):
        appendix = 1
        while True:
            
            newFileName = f'{title}({str(appendix).zfill(2)})'
            newFileName_and_path = path.join(basePath, f'{newFileName}.docx')
            if path.exists(newFileName_and_path):
                appendix += 1
            else:
                docx_filename_path = newFileName_and_path
                break

    # Write the document
    print(f'Saving file to: {docx_filename_path}')
    doc.save(docx_filename_path)
